"""
Generates the ArT Implementation Plan Word document.
Run: python3 generate_doc.py
Output: ProjectRefDocs/ArT_Implementation_Plan.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "ProjectRefDocs", "ArT_Implementation_Plan.docx")

DARK_NAVY  = RGBColor(0x1E, 0x29, 0x3B)
PURPLE     = RGBColor(0x7C, 0x3A, 0xED)
MID_GREY   = RGBColor(0x64, 0x74, 0x8B)


def add_page_numbers(doc):
    """Add centred 'Page X of Y' to the footer of every section."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.font.size = Pt(9)
        run.font.color.rgb = MID_GREY

        def append_field(r, instr):
            fc_begin = OxmlElement("w:fldChar")
            fc_begin.set(qn("w:fldCharType"), "begin")
            it = OxmlElement("w:instrText")
            it.set(qn("xml:space"), "preserve")
            it.text = instr
            fc_end = OxmlElement("w:fldChar")
            fc_end.set(qn("w:fldCharType"), "end")
            r._r.append(fc_begin)
            r._r.append(it)
            r._r.append(fc_end)

        t1 = OxmlElement("w:t")
        t1.set(qn("xml:space"), "preserve")
        t1.text = "Page "
        run._r.append(t1)
        append_field(run, " PAGE ")
        t2 = OxmlElement("w:t")
        t2.set(qn("xml:space"), "preserve")
        t2.text = " of "
        run._r.append(t2)
        append_field(run, " NUMPAGES ")


def add_hyperlinked_toc(doc):
    """Insert a TOC field; Word populates hyperlinks and page numbers on open."""
    para = doc.add_paragraph()
    run  = para.add_run()
    fc_begin = OxmlElement("w:fldChar")
    fc_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fc_sep = OxmlElement("w:fldChar")
    fc_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = (
        "[Open in Word, then right-click here and select "
        "'Update Field' → 'Update entire table' to generate the TOC]"
    )
    fc_end = OxmlElement("w:fldChar")
    fc_end.set(qn("w:fldCharType"), "end")
    run._r.append(fc_begin)
    run._r.append(instr)
    run._r.append(fc_sep)
    run._r.append(placeholder)
    run._r.append(fc_end)
    run.font.color.rgb = MID_GREY
    run.italic = True
    run.font.size = Pt(9)
    doc.add_paragraph()


def set_cell_bg(cell, hex_colour: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_colour)
    tcPr.append(shd)


def add_heading(doc, text, level=1, colour=DARK_NAVY):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = colour
    return h


def add_para(doc, text="", bold=False, italic=False, colour=None, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold   = bold
    r.italic = italic
    if colour:
        r.font.color.rgb = colour
    if size:
        r.font.size = Pt(size)
    return p


def add_bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def set_header_row(row, texts, bg="1E293B", fg=RGBColor(0xFF, 0xFF, 0xFF)):
    for i, cell in enumerate(row.cells):
        cell.text = texts[i] if i < len(texts) else ""
        set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = fg
                run.font.size = Pt(9)


def add_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_header_row(tbl.rows[0], headers)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = tbl.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for row in tbl.rows:
            for i, cell in enumerate(row.cells):
                if i < len(col_widths):
                    cell.width = col_widths[i]
    return tbl


def page_break(doc):
    doc.add_page_break()


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    add_page_numbers(doc)

    # ── TITLE PAGE ────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ArT — Automated Review Tool")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = DARK_NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("FINRA Rule 4210 Pre-Trade Compliance Engine")
    r2.font.size = Pt(16)
    r2.font.color.rgb = PURPLE

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Implementation Plan  |  Version 1.0  |  June 2026").font.color.rgb = MID_GREY
    page_break(doc)

    # ── TABLE OF CONTENTS ─────────────────────────────────────────────────────
    add_heading(doc, "Table of Contents", level=1)
    add_para(doc,
        "After opening in Word, right-click the field below and select "
        "'Update Field' → 'Update entire table' to generate page numbers and hyperlinks.",
        italic=True, colour=MID_GREY, size=9
    )
    add_hyperlinked_toc(doc)
    page_break(doc)

    # ── SECTION 1 ─────────────────────────────────────────────────────────────
    add_heading(doc, "1.  Application / Product Goal", level=1)
    add_para(doc, "Product Name:", bold=True)
    add_para(doc, "ArT — Automated Review Tool")
    doc.add_paragraph()
    add_para(doc, "Purpose:", bold=True)
    add_para(doc,
        "ArT is a real-time pre-trade compliance decision engine designed to enforce FINRA Rule 4210 "
        "(2026 amendments) and Regulation T margin requirements across equity and fixed-income trading "
        "desks. It intercepts every trade before execution, evaluates it against a layered ruleset, "
        "and returns an immediate compliance decision with full reasoning — enabling traders, supervisors, "
        "and the controls team to act on accurate, auditable information."
    )
    doc.add_paragraph()
    add_para(doc, "Primary Goals:", bold=True)
    for g in [
        "Prevent regulatory breaches by hard-blocking trades that violate Reg T initial margin or FINRA 4210 intraday margin requirements.",
        "Support supervisory oversight through an AI-assisted override workflow with risk scoring.",
        "Provide management information (MI) via a dashboard with trend analysis and natural language querying.",
        "Allow compliance teams to author new control rules in plain English without coding.",
        "Maintain an immutable audit trail of every decision, override action, and system event.",
    ]:
        add_bullet(doc, g)
    doc.add_paragraph()
    add_para(doc, "Regulatory Scope:", bold=True)
    add_table(doc,
        ["Regulation", "Scope", "Implemented As"],
        [
            ["Regulation T (12 CFR Part 220)", "Initial margin on equity & fixed-income purchases; short-sale collateral; cash accounts", "Programmatic engine (reg_t_checks.py)"],
            ["FINRA Rule 4210 (2026)", "Intraday Margin Level (IML); 90-day freeze; portfolio margin sub-$5M", "Programmatic engine (finra_4210_checks.py)"],
            ["Firm Custom Rules", "Desk-level controls (analyst notional caps, HY bond limits)", "Database-driven JSON conditions engine"],
        ],
        col_widths=[Inches(2.0), Inches(3.2), Inches(2.2)]
    )
    page_break(doc)

    # ── SECTION 2 ─────────────────────────────────────────────────────────────
    add_heading(doc, "2.  Business Requirements", level=1)
    add_table(doc,
        ["Ref", "Requirement", "Detail"],
        [
            ["BR-01", "Pre-Trade Check", "Every trade must be evaluated by ArT before execution. System must return a decision in real time (<500ms target)."],
            ["BR-02", "Reg T — Long Equity", "Hard-block any equity margin purchase where account equity is below 50% of notional."],
            ["BR-03", "Reg T — Short Sale", "Hard-block short sales where equity is below 150% of notional (100% proceeds + 50% margin)."],
            ["BR-04", "Reg T — Cash Account", "Hard-block cash account purchases where available cash < full notional."],
            ["BR-05", "Reg T — Fixed Income", "Apply product-specific rates: Govt Bond 10%, Municipal 25%, IG Corp 30%, HY Corp 50%."],
            ["BR-06", "FINRA 4210 — IML Safe Harbour", "Soft-block if deficit ≤ min($1,000, 5% of equity). Inform trader of 15-day satisfaction deadline."],
            ["BR-07", "FINRA 4210 — Moderate Deficit", "Soft-block if deficit $1k–$25k. Warn of 90-day freeze risk if deficit_count_30d ≥ 2."],
            ["BR-08", "FINRA 4210 — Large Deficit", "Hard Block with Override if deficit $25k–$100k. Requires supervisory sign-off."],
            ["BR-09", "FINRA 4210 — Critical Deficit", "Hard Block (no override) if deficit > $100k."],
            ["BR-10", "FINRA 4210 — 90-Day Freeze", "Hard-block new short positions when account has active 90-day freeze. Soft-block margin purchases."],
            ["BR-11", "FINRA 4210 — Portfolio Margin", "Soft-block IML-reducing transactions on portfolio margin accounts with equity < $5M."],
            ["BR-12", "Override Workflow", "HARD_BLOCK_WITH_OVERRIDE auto-raises an override request. Supervisors approve, reject, or escalate."],
            ["BR-13", "AI Risk Scoring", "Each override receives a 0–100 AI risk score across LOW / MEDIUM / HIGH / CRITICAL bands."],
            ["BR-14", "Custom Rules", "Controls team can author rules in plain English. AI parser converts them to structured JSON conditions."],
            ["BR-15", "MI Dashboard", "Outcome distribution, daily trend (14d), top rules, top traders, override stats. Period filter: 7/14/30/90d."],
            ["BR-16", "Natural Language Reporting", "Users query decision data in plain English; AI returns a summary and filterable results table."],
            ["BR-17", "Audit Log", "Immutable append-only log of all trade checks, override actions, and system events."],
            ["BR-18", "Multi-Role Support", "Trader, Supervisor, Head of Desk, MD, and Controls Team roles across multiple desks."],
        ],
        col_widths=[Inches(0.7), Inches(1.8), Inches(4.9)]
    )
    page_break(doc)

    # ── SECTION 3 ─────────────────────────────────────────────────────────────
    add_heading(doc, "3.  Design Decisions", level=1)
    add_table(doc,
        ["Ref", "Decision", "Rationale"],
        [
            ["DD-01", "Programmatic engines for Reg T and FINRA 4210",
             "Regulatory thresholds are precise legal requirements and must not be user-editable. Hard-coded Python functions guarantee correctness and auditability."],
            ["DD-02", "Severity escalation: most-severe-wins",
             "When multiple rules fire, the engine returns the worst outcome. Prevents a low-severity warning from masking a critical block. Order: CLEAN < SOFT < OVERRIDE < HARD."],
            ["DD-03", "Automatic override request creation",
             "Override requests are auto-created on HARD_BLOCK_WITH_OVERRIDE. Prevents missed overrides and ensures AI risk scoring runs immediately."],
            ["DD-04", "Simulated AI layer (no external LLM)",
             "NL Rules Builder, Risk Scorer, and NL Reporter are deterministic Python algorithms using regex and heuristics. Avoids LLM API costs and latency in a prototype while retaining the structural design for future LLM integration."],
            ["DD-05", "SQLite for persistence",
             "Eliminates infrastructure dependencies for local development. SQLAlchemy ORM is database-agnostic; migrating to PostgreSQL requires only a connection string change."],
            ["DD-06", "Account state submitted by trader, not looked up",
             "Trade screen requires manual entry of equity, debit balance, IML etc. Deliberate prototype simplification; production would fetch from a position-keeping system."],
            ["DD-07", "React + Vite frontend, FastAPI backend",
             "Clean REST API separation allows independent evolution. FastAPI provides auto OpenAPI docs. Vite provides fast hot-reload."],
            ["DD-08", "Custom rules evaluated by priority before regulatory engines",
             "Firm-level controls are checked first (lower priority number = higher precedence). Regulatory engines then run regardless."],
            ["DD-09", "Audit log is append-only",
             "No UPDATE/DELETE surface on AuditLog. All events are inserts, providing an immutable record for regulatory review."],
            ["DD-10", "days_since_last_deficit stored but not exposed in UI",
             "Field is persisted on the Trade record for future use in escalation logic (e.g. auto-triggering freeze) but the prototype does not yet collect or use it."],
        ],
        col_widths=[Inches(0.7), Inches(2.0), Inches(4.7)]
    )
    page_break(doc)

    # ── SECTION 4 ─────────────────────────────────────────────────────────────
    add_heading(doc, "4.  Architecture Details", level=1)

    add_heading(doc, "4.1  Stack Summary", level=2)
    add_table(doc,
        ["Layer", "Technology", "Version", "Role"],
        [
            ["Frontend",      "React + TypeScript",  "React 18",  "Single-page application; all UI screens"],
            ["Build Tool",    "Vite",                "5.x",       "Dev server (port 5173), production bundler"],
            ["Styling",       "Tailwind CSS",        "3.x",       "Utility-first CSS"],
            ["Charts",        "Recharts",            "2.x",       "Bar, pie, trend charts on MI Dashboard"],
            ["Backend",       "FastAPI",             "0.111",     "REST API; auto OpenAPI docs at /docs"],
            ["ORM",           "SQLAlchemy",          "2.0",       "Database abstraction layer"],
            ["Database",      "SQLite",              "—",         "File-based; art.db in backend root"],
            ["Migrations",    "Alembic",             "1.13",      "Schema version control"],
            ["Runtime",       "Python",              "3.9",       "Backend language"],
        ],
        col_widths=[Inches(1.1), Inches(1.5), Inches(0.9), Inches(3.9)]
    )

    add_heading(doc, "4.2  Directory Structure", level=2)
    p = doc.add_paragraph()
    r = p.add_run(
        "finra4210/\n"
        "├── backend/\n"
        "│   ├── app/\n"
        "│   │   ├── engine/\n"
        "│   │   │   ├── rules_engine.py        ← Orchestrator\n"
        "│   │   │   ├── reg_t_checks.py        ← Regulation T checks\n"
        "│   │   │   ├── finra_4210_checks.py   ← FINRA 4210 checks\n"
        "│   │   │   └── ai_simulator.py        ← NL Builder / Risk Scorer / NL Reporter\n"
        "│   │   ├── models/                    ← SQLAlchemy ORM models\n"
        "│   │   ├── routers/                   ← FastAPI route handlers\n"
        "│   │   │   └── art.py                 ← POST /art/check (primary endpoint)\n"
        "│   │   ├── schemas/                   ← Pydantic request/response models\n"
        "│   │   ├── main.py                    ← App factory, router registration\n"
        "│   │   └── database.py                ← SQLAlchemy engine + session\n"
        "│   ├── seed_data.py                   ← Seeds users + rules\n"
        "│   └── requirements.txt\n"
        "└── frontend/\n"
        "    └── src/\n"
        "        ├── pages/\n"
        "        │   ├── TraderPage.tsx          ← Trade entry + decision panel\n"
        "        │   ├── DashboardPage.tsx       ← MI Dashboard + NL Query\n"
        "        │   ├── ApprovalsPage.tsx       ← Supervisor override workflow\n"
        "        │   ├── RulesPage.tsx           ← Rules Manager + NL Builder\n"
        "        │   └── AuditPage.tsx           ← Audit log viewer\n"
        "        ├── services/api.ts             ← All API calls\n"
        "        └── types/index.ts              ← TypeScript interfaces"
    )
    r.font.name = "Courier New"
    r.font.size = Pt(8)

    add_heading(doc, "4.3  Request Flow — Pre-Trade Check", level=2)
    add_table(doc,
        ["Step", "Action", "Detail"],
        [
            ["1",  "Trade submitted",               "POST /art/check with TradeCreate payload"],
            ["2",  "Persist Trade record",          "Status = PENDING_CHECK; notional computed server-side"],
            ["3",  "rules_engine.run_engine()",     "Receives trade_data dict + trader_data dict"],
            ["4",  "Pass 1: Custom rules",          "DB rules loaded, filtered by asset class, evaluated in priority order"],
            ["5",  "Pass 2: Regulation T",          "check_reg_t() — initial margin, short collateral, cash account"],
            ["6",  "Pass 3: FINRA 4210",            "check_finra_4210() — IML deficit bands, 90-day freeze, portfolio margin"],
            ["7",  "Worst outcome determined",      "Severity ranking applied across all triggered rules"],
            ["8",  "ArtDecision persisted",         "Outcome, rules_triggered (JSON), reasoning, margin figures, processing time"],
            ["9",  "Override path?",                "If HARD_BLOCK_WITH_OVERRIDE → AI risk score computed → OverrideRequest created"],
            ["10", "AuditLog written",              "TRADE_CHECKED event with metadata (immutable)"],
            ["11", "Response returned",             "TradeCheckResponse: trade + decision + override_request_id"],
        ],
        col_widths=[Inches(0.5), Inches(2.0), Inches(4.9)]
    )

    add_heading(doc, "4.4  API Endpoints", level=2)
    add_table(doc,
        ["Method", "Path", "Description"],
        [
            ["POST",  "/art/check",               "Submit trade for compliance check (primary endpoint)"],
            ["GET",   "/art/decisions",            "List decisions, filterable by outcome"],
            ["GET",   "/overrides/",               "List override requests, filterable by status"],
            ["POST",  "/overrides/{id}/action",    "Approve or reject an override"],
            ["POST",  "/overrides/{id}/escalate",  "Escalate override to senior approver"],
            ["GET",   "/rules/",                   "List all rules"],
            ["POST",  "/rules/parse-nl",           "AI: parse plain English into rule structure"],
            ["POST",  "/rules/from-nl",            "AI: create rule from plain English"],
            ["PATCH", "/rules/{id}",               "Update rule (e.g. toggle is_active)"],
            ["GET",   "/reports/summary",          "Dashboard KPI summary"],
            ["GET",   "/reports/daily-trend",      "14-day decision trend"],
            ["GET",   "/reports/top-rules",        "Top triggered rules"],
            ["GET",   "/reports/top-traders",      "Traders with most block decisions"],
            ["GET",   "/reports/override-stats",   "Override approval statistics"],
            ["POST",  "/reports/nl-query",         "AI: natural language query"],
            ["GET",   "/audit/",                   "Audit log, filterable by event type and date range"],
            ["GET",   "/users/",                   "List users"],
        ],
        col_widths=[Inches(0.7), Inches(2.5), Inches(4.2)]
    )
    page_break(doc)

    # ── SECTION 5 ─────────────────────────────────────────────────────────────
    add_heading(doc, "5.  Data Dictionary", level=1)

    add_heading(doc, "5a.  UI-Exposed Fields — Trader Manual Entry Screen", level=2)
    add_para(doc, "All fields below are entered by the trader and submitted to the ArT engine via POST /art/check.")
    add_table(doc,
        ["UI Label", "Field Name", "Type / Values", "Required?", "Description"],
        [
            ["Trader",                "trader_id",                    "Integer (FK → users)",                                                      "Yes", "Trader submitting the trade. Drives seniority-based custom rule conditions."],
            ["Account Type",          "account_type",                 "CASH / MARGIN / PORTFOLIO_MARGIN",                                          "Yes", "Determines which margin rules apply. CASH = full payment required."],
            ["Asset Class",           "asset_class",                  "EQUITY / FIXED_INCOME",                                                     "Yes", "Drives Reg T rates and rule scope filtering."],
            ["Product Type",          "product_type",                 "STOCK, ETF, EQUITY_OPTION, GOVT_BOND, CORP_BOND_IG, CORP_BOND_HY, MUNICIPAL_BOND", "Yes", "Specific Reg T initial margin rate (10%–50%)."],
            ["Direction",             "direction",                    "BUY / SELL / SHORT_SELL / BUY_TO_COVER / SHORT_PUT / SHORT_CALL",           "Yes", "SHORT_SELL / SHORT_PUT / SHORT_CALL are IML-reducing transactions (Rule 4210(a)(18))."],
            ["Ticker",                "ticker",                       "String (max 20)",                                                           "Yes", "Instrument identifier. Stored for audit; not validated against a security master."],
            ["Quantity",              "quantity",                     "Float",                                                                     "Yes", "Number of units. Combined with price to compute notional server-side."],
            ["Price ($)",             "price",                        "Float",                                                                     "Yes", "Unit price. Notional = quantity × price (computed server-side, not accepted from client)."],
            ["Credit Rating",         "credit_rating",                "String (max 10), nullable",                                                 "No",  "Fixed Income only. Used to distinguish IG vs HY where product type is ambiguous. e.g. BBB+."],
            ["Account Equity ($)",    "account_equity",               "Float",                                                                     "Yes", "Total net worth of the account. Primary figure for all margin checks."],
            ["Existing Long MV ($)",  "existing_long_market_value",   "Float, default 0",                                                          "No",  "Market value of existing long positions. Used in post-trade maintenance margin calculation."],
            ["Existing Short MV ($)", "existing_short_market_value",  "Float, default 0",                                                          "No",  "Market value of existing short positions. Each unit carries 30% maintenance margin."],
            ["Debit Balance ($)",     "existing_debit_balance",       "Float, default 0",                                                          "No",  "Margin loan outstanding. Subtracted from equity to derive net available margin."],
            ["Intraday Margin Level ($)", "intraday_margin_level",    "Float, default 0",                                                          "No",  "Current IML snapshot. If 0, engine computes IML from equity/LMV/SMV/debit. If supplied, used as pre-trade starting point."],
            ["Deficit Count (30d)",   "deficit_count_30d",            "Integer, default 0",                                                        "No",  "Intraday margin deficits in past 30 days. Escalates warning language; ≥2 flags freeze risk."],
            ["90-Day Freeze Active",  "is_90day_freeze_active",       "Boolean (stored as Integer 0/1)",                                           "No",  "When ON: hard-blocks all new short positions (F4210-FRZ-001). Soft-blocks margin purchases (F4210-FRZ-002)."],
        ],
        col_widths=[Inches(1.2), Inches(1.5), Inches(1.8), Inches(0.7), Inches(3.2)]
    )

    add_heading(doc, "5b.  Backend / Database-Only Fields", level=2)
    add_para(doc, "Generated by the system and stored in the database. Not entered by users.", italic=True, colour=MID_GREY)

    doc.add_heading("Table: trades", level=3)
    add_table(doc, ["Field", "Type", "Description"],
        [
            ["id",                     "Integer PK",      "Auto-increment primary key."],
            ["trade_ref",              "String (20)",     "System-generated unique ref e.g. TRD-ABCD1234."],
            ["notional",               "Float",           "Computed server-side as quantity × price. Never accepted from client."],
            ["status",                 "Enum",            "PENDING_CHECK → CLEAN_PASS / SOFT_BLOCK / HARD_BLOCK / HARD_BLOCK_WITH_OVERRIDE / OVERRIDE_APPROVED / OVERRIDE_REJECTED."],
            ["trade_date",             "DateTime (UTC)",  "Timestamp when trade was created."],
            ["created_at",             "DateTime (UTC)",  "Timestamp of record insertion."],
            ["days_since_last_deficit","Integer, null",   "Days since the account last had an unresolved deficit. Stored for future escalation logic; not yet collected via UI."],
        ],
        col_widths=[Inches(1.8), Inches(1.3), Inches(4.3)]
    )

    doc.add_heading("Table: art_decisions", level=3)
    add_table(doc, ["Field", "Type", "Description"],
        [
            ["id",                      "Integer PK",     "Auto-increment primary key."],
            ["decision_ref",            "String (20)",    "System-generated ref e.g. DEC-ABCD1234."],
            ["trade_id",                "Integer FK",     "Links to trades.id (1:1 relationship)."],
            ["outcome",                 "Enum",           "CLEAN_PASS / SOFT_BLOCK / HARD_BLOCK / HARD_BLOCK_WITH_OVERRIDE."],
            ["rules_triggered",         "JSON array",     "Full list of violation objects: rule_code, rule_name, block_type, description, margin figures."],
            ["reasoning",               "Text",           "Human-readable ArT summary returned to the trader UI."],
            ["reg_t_margin_required",   "Float, null",    "Required margin for first Reg T violation (if any)."],
            ["reg_t_margin_available",  "Float, null",    "Available margin at time of check."],
            ["intraday_margin_deficit", "Float, null",    "Size of IML deficit created (if any FINRA 4210 IML rule fired)."],
            ["processing_time_ms",      "Integer",        "Engine execution time in milliseconds."],
            ["created_at",              "DateTime (UTC)", "Timestamp of record insertion."],
        ],
        col_widths=[Inches(2.0), Inches(1.2), Inches(4.2)]
    )

    doc.add_heading("Table: override_requests", level=3)
    add_table(doc, ["Field", "Type", "Description"],
        [
            ["id",                   "Integer PK",      "Auto-increment primary key."],
            ["override_ref",         "String (20)",     "System-generated ref e.g. OVR-ABCD1234."],
            ["decision_id",          "Integer FK",      "Links to art_decisions.id (1:1)."],
            ["trade_id",             "Integer FK",      "Links to trades.id."],
            ["requested_by_id",      "Integer FK",      "User ID of the trader — auto-set from the trade record."],
            ["approver_id",          "Integer FK, null","User ID of the supervisor who actioned the request."],
            ["status",               "Enum",            "PENDING / APPROVED / REJECTED / ESCALATED."],
            ["ai_risk_score",        "Float (0–100)",   "AI-generated risk score. Factors: notional, deficit ratio, trader seniority, override history, Reg T presence."],
            ["ai_risk_band",         "Enum",            "LOW / MEDIUM / HIGH / CRITICAL."],
            ["ai_risk_rationale",    "Text",            "Full AI explanation of risk score factors and recommendation."],
            ["trader_justification", "Text, null",      "Trader's written justification. Not yet collected via UI; reserved for future enhancement."],
            ["approver_notes",       "Text, null",      "Notes entered by supervisor on approval/rejection."],
            ["requested_at",         "DateTime (UTC)",  "Timestamp when override was auto-created."],
            ["resolved_at",          "DateTime, null",  "Timestamp when override was actioned."],
        ],
        col_widths=[Inches(1.8), Inches(1.3), Inches(4.3)]
    )

    doc.add_heading("Table: rules", level=3)
    add_table(doc, ["Field", "Type", "Description"],
        [
            ["id",                          "Integer PK",   "Auto-increment primary key."],
            ["rule_code",                   "String (20)",  "Unique code e.g. F4210-FRZ-001, REGT-LM-001, CUST-001."],
            ["name",                        "String (200)", "Short display name."],
            ["description",                 "Text",         "Technical description."],
            ["natural_language_description","Text, null",   "Plain English description shown in Rules Manager and NL Builder output."],
            ["rule_type",                   "Enum",         "REG_T / FINRA_4210 / CUSTOM."],
            ["asset_class_scope",           "Enum",         "EQUITY / FIXED_INCOME / ALL. Engine filters rules by trade asset class."],
            ["block_type",                  "Enum",         "SOFT_BLOCK / HARD_BLOCK / HARD_BLOCK_WITH_OVERRIDE."],
            ["conditions",                  "JSON, null",   "Conditions tree (AND/OR groups). Only evaluated for CUSTOM rules; REG_T and FINRA_4210 use programmatic engines."],
            ["is_active",                   "Boolean",      "If false, excluded from engine evaluation."],
            ["priority",                    "Integer",      "Lower = evaluated first. Custom rules run before regulatory engines."],
            ["created_by",                  "String",       "SYSTEM for seed rules; AI_NL_BUILDER for NL-created rules."],
            ["created_at",                  "DateTime",     "UTC timestamp."],
            ["updated_at",                  "DateTime",     "Updated on any PATCH."],
        ],
        col_widths=[Inches(2.0), Inches(1.2), Inches(4.2)]
    )

    doc.add_heading("Table: users", level=3)
    add_table(doc, ["Field", "Type", "Description"],
        [
            ["id",              "Integer PK",   "Auto-increment primary key."],
            ["username",        "String (50)",  "Unique handle e.g. jsmith."],
            ["full_name",       "String (100)", "Display name."],
            ["role",            "Enum",         "TRADER / SUPERVISOR / HEAD_OF_DESK / MANAGING_DIRECTOR / CONTROLS_TEAM / ADMIN."],
            ["seniority_level", "Enum",         "ANALYST / ASSOCIATE / VP / DIRECTOR / MANAGING_DIRECTOR. Used in CUST-001 condition."],
            ["desk",            "String (50)",  "Trading desk e.g. Equity Derivatives, Fixed Income."],
            ["is_active",       "Boolean",      "Soft-delete flag."],
            ["created_at",      "DateTime",     "UTC timestamp."],
        ],
        col_widths=[Inches(1.5), Inches(1.2), Inches(4.7)]
    )

    doc.add_heading("Table: audit_log", level=3)
    add_table(doc, ["Field", "Type", "Description"],
        [
            ["id",             "Integer PK",   "Auto-increment primary key."],
            ["event_type",     "String (50)",  "TRADE_CHECKED / OVERRIDE_APPROVED / OVERRIDE_REJECTED / OVERRIDE_ESCALATED / OVERRIDE_PENDING."],
            ["entity_type",    "String (50)",  "TRADE / OVERRIDE."],
            ["entity_id",      "Integer, null","PK of the related entity."],
            ["user_id",        "Integer, null","ID of the acting user."],
            ["username",       "String (50)",  "Denormalised username for fast query without join."],
            ["description",    "Text",         "Human-readable event description."],
            ["event_metadata", "JSON, null",   "Structured event data e.g. trade_ref, outcome, notional, rules_count."],
            ["created_at",     "DateTime",     "UTC timestamp. Indexed for date-range queries."],
        ],
        col_widths=[Inches(1.5), Inches(1.5), Inches(4.4)]
    )
    page_break(doc)

    # ── SECTION 6 ─────────────────────────────────────────────────────────────
    add_heading(doc, "6.  Control Logic — Rules Engine Flow", level=1)
    add_para(doc,
        "The following describes the decision logic executed by rules_engine.run_engine() on every "
        "trade submission. Three sequential passes run; the worst outcome across all passes is returned."
    )

    add_heading(doc, "6.1  High-Level Engine Flow", level=2)
    add_table(doc, ["Stage", "Action", "Detail"],
        [
            ["START",  "Trade received",            "POST /art/check — TradeCreate payload validated by Pydantic"],
            ["1",      "Persist Trade",             "Trade record saved with status PENDING_CHECK. Notional computed."],
            ["2",      "Pass 1 — Custom Rules",     "Load active rules from DB filtered by asset class. Evaluate JSON conditions tree in priority order."],
            ["3",      "Pass 2 — Regulation T",     "check_reg_t(): initial margin (long), short sale collateral, cash account full payment, maintenance margin."],
            ["4",      "Pass 3 — FINRA 4210",       "check_finra_4210(): 90-day freeze check, IML deficit bands, portfolio margin threshold."],
            ["5",      "Determine worst outcome",   "Severity: CLEAN(0) < SOFT(1) < OVERRIDE(2) < HARD(3). Worst across all triggered rules wins."],
            ["6",      "Build reasoning text",      "Summarise triggered rules and outcome into human-readable string for trader UI."],
            ["7",      "Persist ArtDecision",       "Save outcome, rules_triggered JSON, reasoning, margin figures, processing_time_ms."],
            ["8",      "Override branch",           "If HARD_BLOCK_WITH_OVERRIDE: score_override_risk() → create OverrideRequest with AI score + rationale."],
            ["9",      "Write AuditLog",            "Append TRADE_CHECKED event. Immutable — no update/delete path."],
            ["END",    "Return response",           "TradeCheckResponse: trade + decision + override_request_id (null if no override)."],
        ],
        col_widths=[Inches(0.7), Inches(2.0), Inches(4.7)]
    )

    add_heading(doc, "6.2  FINRA 4210 IML Decision Tree", level=2)
    add_table(doc, ["Condition", "Branch", "Rule Code & Outcome"],
        [
            ["Is 90-day freeze active?",           "YES — direction = SHORT_SELL / SHORT_PUT / SHORT_CALL", "F4210-FRZ-001 → HARD_BLOCK"],
            ["Is 90-day freeze active?",           "YES — direction = BUY",                                 "F4210-FRZ-002 → SOFT_BLOCK"],
            ["Is direction IML-reducing?",         "NO (SELL / BUY_TO_COVER)",                              "No FINRA 4210 IML check. Exit."],
            ["Compute IML after trade",            "IML_after = IML_before – (notional × maint. rate)",     "25% for BUY; 30% for short directions."],
            ["IML_after ≥ 0?",                     "YES",                                                   "No FINRA 4210 IML violation."],
            ["Deficit ≤ min($1,000, 5% equity)?",  "YES",                                                   "F4210-IML-001 → SOFT_BLOCK (safe harbour warning)"],
            ["Deficit < $25,000?",                 "YES",                                                   "F4210-IML-002 → SOFT_BLOCK (moderate deficit)"],
            ["Deficit < $100,000?",                "YES",                                                   "F4210-IML-003 → HARD_BLOCK_WITH_OVERRIDE"],
            ["Deficit ≥ $100,000",                 "—",                                                     "F4210-IML-004 → HARD_BLOCK (no override path)"],
            ["Portfolio margin + equity < $5M?",   "YES + direction IML-reducing",                          "F4210-PM-001 → SOFT_BLOCK"],
        ],
        col_widths=[Inches(2.5), Inches(2.2), Inches(2.7)]
    )

    add_heading(doc, "6.3  Regulation T Decision Tree", level=2)
    add_table(doc, ["Condition", "Branch", "Rule Code & Outcome"],
        [
            ["Account type = CASH?",                        "YES — direction = BUY / BUY_TO_COVER and equity < notional",   "REGT-CA-001 → HARD_BLOCK. No further Reg T checks."],
            ["Direction = BUY (margin account)?",           "margin_available ≥ margin_required",                           "No Reg T initial margin violation. Check maintenance."],
            ["Direction = BUY — shortfall?",                "Coverage ratio ≥ 80% of required rate",                        "REGT-LM-001 → HARD_BLOCK_WITH_OVERRIDE"],
            ["Direction = BUY — shortfall?",                "Coverage ratio < 80% of required rate",                        "REGT-LM-001 → HARD_BLOCK"],
            ["Direction = SHORT_SELL / SHORT_PUT / SHORT_CALL?", "Required collateral = notional × 1.50",                  "—"],
            ["Short sale — shortfall?",                     "Coverage ≥ 85% of required collateral",                        "REGT-SS-001 → HARD_BLOCK_WITH_OVERRIDE"],
            ["Short sale — shortfall?",                     "Coverage < 85% of required collateral",                        "REGT-SS-001 → HARD_BLOCK"],
            ["Post-trade net equity < 25% of total LMV?",  "YES",                                                           "REGT-MM-001 → SOFT_BLOCK (maintenance warning)"],
        ],
        col_widths=[Inches(2.5), Inches(2.2), Inches(2.7)]
    )

    add_heading(doc, "6.4  AI Override Risk Score Factors", level=2)
    add_table(doc, ["Factor", "Max Points", "Logic"],
        [
            ["Trade notional size",    "30",  "Log₁₀ scale × 5. $1M ≈ 15pts, $10M ≈ 20pts, $100M ≈ 25pts."],
            ["IML deficit severity",   "25",  "(deficit / account_equity) × 100, capped at 25."],
            ["Trader seniority",       "25",  "Analyst=25, Associate=18, VP=10, Director=5, MD=2."],
            ["Override history (30d)", "15",  "(approved_30d × 3) + (rejected_30d × 2), capped at 15."],
            ["Reg T violation present","10",  "Flat +10 if any REGT-prefixed rule triggered."],
            ["Total / Bands",          "100", "≤30 = LOW · 31–60 = MEDIUM · 61–80 = HIGH · >80 = CRITICAL"],
        ],
        col_widths=[Inches(2.0), Inches(1.0), Inches(4.4)]
    )
    page_break(doc)

    # ── APPENDIX A ────────────────────────────────────────────────────────────
    add_heading(doc, "Appendix A  –  All User Prompts", level=1)
    add_para(doc, "Every prompt submitted during the development and review session.", italic=True, colour=MID_GREY)
    add_table(doc,
        ["Ref", "Category", "Prompt Text"],
        [
            ["P-01", "Session setup",      "cd backend && pip install -r requirements.txt && python seed_data.py && uvicorn app.main:app --reload"],
            ["P-02", "Navigation",         "pwd"],
            ["P-03", "Project location",   "myproj1"],
            ["P-04", "Navigation",         "cd ~/Documents"],
            ["P-05", "Shell alias query",  "In a different Terminal window I created an alias for the path /Users/athithi/Documents/AI_Study/finra4210/. Should I create it every time when a new terminal is opened?"],
            ["P-06", "Alias name",         "finra is good"],
            ["P-07", "PATH fix",           "sure (accepted Claude's offer to add Python 3.9 bin to PATH in ~/.zshrc)"],
            ["P-08", "Frontend URL",       "From the other terminal I can see that the React frontend service is running in http://localhost:5173"],
            ["P-09", "App launch",         "How do I launch the app?"],
            ["P-10", "Query method",       "I have some queries regarding the app as I am reviewing different features. Please can you advise me what is the best way to post those queries to you? Should I ask one by one here in the command prompt or should I compile an initial list and then attach them here?"],
            ["P-11", "Read queries file",  "Ok. Here is the path to my queries: 'finra/ProjectRefDocs/Queries.md'. Please answer them."],
            ["P-12", "Enrich test data",   "Can you enrich the table in your answer to my Q2 with test data for Account State"],
            ["P-13", "Field explanation",  "Can you explain about the 6 fields in the Account State section of the manual entry screen?"],
            ["P-14", "Documentation",      "Ask 1 — Can you generate the following documentation for the ArT application in downloadable word documents: Implementation Plan with product goal, business requirements, design decisions, architecture, data dictionary, control logic flowchart, Appendix of prompts and Claude decision points."],
        ],
        col_widths=[Inches(0.6), Inches(1.4), Inches(5.4)]
    )
    page_break(doc)

    # ── APPENDIX B ────────────────────────────────────────────────────────────
    add_heading(doc, "Appendix B  –  Claude Decision Points & Queries Needing User Input", level=1)
    add_para(doc, "Every point where Claude asked the user a question or required input to proceed.", italic=True, colour=MID_GREY)
    add_table(doc,
        ["Ref", "Decision Point", "Context", "User Response / Resolution"],
        [
            ["D-01", "Project location unknown",
             "Running 'cd backend' from ~/ failed — no such directory. Claude asked: 'Could you point me to where your project lives?'",
             "User replied 'myproj1' (not found), then 'cd ~/Documents'. Path finally located at AI_Study/finra4210/."],
            ["D-02", "pip not found on PATH",
             "pip command not found on macOS Python 3.9. Claude switched to pip3 automatically and informed the user.",
             "Resolved autonomously — no user input required."],
            ["D-03", "Port 8000 already in use",
             "Uvicorn start failed with 'Address already in use'. Claude informed the user rather than killing the existing process.",
             "User confirmed the server in the other terminal was sufficient. No action needed."],
            ["D-04", "Python scripts not on PATH",
             "pip3 install produced multiple PATH warnings for uvicorn, alembic, faker etc. Claude offered to add the Python 3.9 bin directory to ~/.zshrc.",
             "User replied 'sure'. Claude added the PATH export and verified with 'which uvicorn'."],
            ["D-05", "Alias name confirmation",
             "After explaining how to make an alias permanent, Claude asked: 'What was the alias name you used? I can tailor the command exactly.'",
             "User replied 'finra is good'. Claude used 'finra' and added it to ~/.zshrc."],
            ["D-06", "Query submission method",
             "User asked whether to submit queries one at a time or compile a file. Claude recommended compiling a file.",
             "User provided path: finra/ProjectRefDocs/Queries.md."],
            ["D-07", "Queries file read",
             "Claude read Queries.md and found 3 questions (F4210-FRZ-001 freeze accounts; counterparty list for breach testing; trade detail drill-down). All answered in one response.",
             "No further input required."],
            ["D-08", "Word document generation — output location",
             "User requested 'downloadable word documents'. Claude chose ProjectRefDocs/ as output location based on context, without asking.",
             "No explicit confirmation requested. Claude proceeded with python-docx and saved to ProjectRefDocs/ArT_Implementation_Plan.docx."],
        ],
        col_widths=[Inches(0.5), Inches(1.6), Inches(2.8), Inches(2.5)]
    )

    # ── SAVE ──────────────────────────────────────────────────────────────────
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
