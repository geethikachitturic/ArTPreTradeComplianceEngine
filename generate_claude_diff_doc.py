"""
Generates: ReadingMaterials/Diff between Claude Code and Claude ai web.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.expanduser(
    "~/Documents/AI_Study/ReadingMaterials/Diff between Claude Code and Claude ai web.docx"
)

DARK_NAVY = RGBColor(0x1E, 0x29, 0x3B)
PURPLE    = RGBColor(0x7C, 0x3A, 0xED)
MID_GREY  = RGBColor(0x64, 0x74, 0x8B)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_bg(cell, hex_colour):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_colour)
    tcPr.append(shd)


def add_heading(doc, text, level=1, colour=DARK_NAVY):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = colour
    return h


def add_para(doc, text="", bold=False, italic=False, colour=None, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold   = bold
    r.italic = italic
    if colour:
        r.font.color.rgb = colour
    if size:
        r.font.size = Pt(size)
    return p


def add_bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def set_header_row(row, texts, bg="1E293B"):
    for i, cell in enumerate(row.cells):
        cell.text = texts[i] if i < len(texts) else ""
        set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(9)


def add_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_header_row(tbl.rows[0], headers)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = tbl.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for row in tbl.rows:
            for i, cell in enumerate(row.cells):
                if i < len(col_widths):
                    cell.width = col_widths[i]
    return tbl


def page_break(doc):
    doc.add_page_break()


def add_page_numbers(doc):
    """Add centred 'Page X of Y' to the footer of every section."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.font.size = Pt(9)
        run.font.color.rgb = MID_GREY

        def append_field(r, instr):
            fc_begin = OxmlElement("w:fldChar")
            fc_begin.set(qn("w:fldCharType"), "begin")
            it = OxmlElement("w:instrText")
            it.set(qn("xml:space"), "preserve")
            it.text = instr
            fc_end = OxmlElement("w:fldChar")
            fc_end.set(qn("w:fldCharType"), "end")
            r._r.append(fc_begin)
            r._r.append(it)
            r._r.append(fc_end)

        t1 = OxmlElement("w:t")
        t1.set(qn("xml:space"), "preserve")
        t1.text = "Page "
        run._r.append(t1)

        append_field(run, " PAGE ")

        t2 = OxmlElement("w:t")
        t2.set(qn("xml:space"), "preserve")
        t2.text = " of "
        run._r.append(t2)

        append_field(run, " NUMPAGES ")


def add_hyperlinked_toc(doc):
    """
    Insert a TOC field. Word populates page numbers and hyperlinks
    when the file is opened (right-click → Update Field).
    """
    para = doc.add_paragraph()
    run  = para.add_run()

    fc_begin = OxmlElement("w:fldChar")
    fc_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '

    fc_sep = OxmlElement("w:fldChar")
    fc_sep.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = (
        "[Open in Word, then right-click here and select "
        "'Update Field' → 'Update entire table' to generate the TOC]"
    )

    fc_end = OxmlElement("w:fldChar")
    fc_end.set(qn("w:fldCharType"), "end")

    run._r.append(fc_begin)
    run._r.append(instr)
    run._r.append(fc_sep)
    run._r.append(placeholder)
    run._r.append(fc_end)

    run.font.color.rgb = MID_GREY
    run.italic = True
    run.font.size = Pt(9)
    doc.add_paragraph()


def build():
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    add_page_numbers(doc)

    # ── TITLE PAGE ────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Claude Code vs Claude.ai Web")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = DARK_NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Understanding the Difference Between the Two Products")
    r2.font.size = Pt(14)
    r2.font.color.rgb = PURPLE

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Reference Guide  |  June 2026").font.color.rgb = MID_GREY
    page_break(doc)

    # ── TABLE OF CONTENTS ─────────────────────────────────────────────────────
    add_heading(doc, "Table of Contents", level=1)
    add_para(doc,
        "After opening in Word, right-click the field below and select "
        "'Update Field' → 'Update entire table' to generate page numbers and hyperlinks.",
        italic=True, colour=MID_GREY, size=9
    )
    add_hyperlinked_toc(doc)
    page_break(doc)

    # ── SECTION 1 ─────────────────────────────────────────────────────────────
    add_heading(doc, "1.  Overview", level=1)
    add_para(doc,
        "Anthropic provides two distinct ways to interact with Claude. Understanding the difference "
        "is important for choosing the right tool and knowing how context, memory, and instructions "
        "are managed in each."
    )
    doc.add_paragraph()
    add_table(doc,
        ["Dimension", "Claude Code (CLI)", "Claude.ai Web UI"],
        [
            ["What it is",        "A command-line tool run in a terminal",                   "A browser-based chat interface at claude.ai"],
            ["Primary use",       "Software engineering, file editing, running commands",    "General queries, writing, analysis, chat"],
            ["How you launch it", "Terminal: type 'claude' from any directory",              "Browser: navigate to claude.ai"],
            ["Context source",    "Current directory files, CLAUDE.md, memory system",      "Conversation history + Project knowledge files"],
            ["Instructions",      "CLAUDE.md file in the project or home directory",         "Instructions section inside a Project"],
            ["Memory",            "File-based memory at ~/.claude/projects/",               "Conversation history saved per Project"],
            ["File access",       "Full read/write access to your local filesystem",         "No direct filesystem access"],
            ["Code execution",    "Can run shell commands, scripts, tests",                  "Cannot execute code on your machine"],
        ],
        col_widths=[Inches(1.5), Inches(2.5), Inches(2.5)]
    )
    page_break(doc)

    # ── SECTION 2 ─────────────────────────────────────────────────────────────
    add_heading(doc, "2.  Claude Code (CLI) — How It Works", level=1)
    add_para(doc,
        "Claude Code is the tool used in the ArT project sessions. "
        "It runs locally in your terminal and has direct access to your files and system."
    )

    add_heading(doc, "2.1  Context Sources", level=2)
    add_table(doc,
        ["Source", "Location", "Purpose"],
        [
            ["CLAUDE.md",       "Project root (e.g. finra4210/CLAUDE.md) or ~/.claude/CLAUDE.md for global",  "Persistent instructions applied every time Claude Code launches from that directory"],
            ["Memory files",    "~/.claude/projects/<project-path>/memory/",                                   "Notes saved across sessions — preferences, project context, working style feedback"],
            ["Live file reads", "Any file in the working directory",                                            "Claude reads files on demand during the session — code, config, docs, data"],
        ],
        col_widths=[Inches(1.4), Inches(2.8), Inches(2.3)]
    )

    add_heading(doc, "2.2  Directory Structure", level=2)
    p = doc.add_paragraph()
    r = p.add_run(
        "~/.claude/\n"
        "├── settings.json                          ← Global Claude Code settings\n"
        "├── history.jsonl                           ← Shell command history\n"
        "├── CLAUDE.md                               ← Global instructions (applies everywhere)\n"
        "└── projects/\n"
        "    └── -Users-athithi-Documents-AI_Study-finra4210/\n"
        "        ├── memory/\n"
        "        │   ├── MEMORY.md                  ← Index of saved memories\n"
        "        │   ├── project_setup.md            ← Notes on ArT tech stack & scope\n"
        "        │   └── user_profile.md             ← Notes on user as Product Owner\n"
        "        └── sessions/                       ← Conversation session data\n\n"
        "~/Documents/AI_Study/finra4210/\n"
        "└── CLAUDE.md                               ← Project-level instructions"
    )
    r.font.name = "Courier New"
    r.font.size = Pt(8)
    r.font.color.rgb = DARK_NAVY

    add_heading(doc, "2.3  Memory System", level=2)
    add_para(doc,
        "When Claude Code learns something useful about you or the project, it saves it to a memory "
        "file. These files are loaded automatically at the start of future sessions, so Claude already knows:"
    )
    for item in [
        "You are the Product Owner for ArT (a FINRA Rule 4210 pre-trade compliance engine)",
        "The project is at ~/Documents/AI_Study/finra4210/",
        "Preferred output is Word (.docx) saved to ProjectRefDocs/",
        "Tech stack: FastAPI backend, React/TypeScript frontend, SQLite, SQLAlchemy ORM",
    ]:
        add_bullet(doc, item)
    page_break(doc)

    # ── SECTION 3 ─────────────────────────────────────────────────────────────
    add_heading(doc, "3.  Claude.ai Web UI Projects — How It Works", level=1)
    add_para(doc,
        "The Projects feature on claude.ai gives you a structured, persistent workspace in the browser. "
        "It is separate from Claude Code and has no access to your local filesystem."
    )
    doc.add_paragraph()
    add_table(doc,
        ["Feature", "Description"],
        [
            ["Instructions section", "A persistent system prompt written once, applied to every conversation in the project. Equivalent to CLAUDE.md in Claude Code."],
            ["Project Knowledge",    "Files and documents you upload (PDFs, text, code). Claude references these across all conversations in the project."],
            ["Conversation History", "All chats within a project are saved and browsable. You can return to any previous conversation."],
            ["No filesystem access", "Claude.ai cannot read or write files on your machine. Content must be pasted or uploaded manually."],
            ["No code execution",    "Claude.ai cannot run scripts, commands, or tests. It generates code for you to run yourself."],
            ["Collaboration",        "Projects can be shared with team members on paid plans."],
        ],
        col_widths=[Inches(1.8), Inches(5.6)]
    )

    add_heading(doc, "3.1  Equivalent Instructions for ArT on Claude.ai", level=2)
    add_para(doc,
        "If you wanted to mirror the ArT working context in a Claude.ai Project, "
        "the Instructions section would contain:"
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(
        "I am the Product Owner for ArT — an Automated Review Tool enforcing\n"
        "FINRA Rule 4210 (2026 amendments) and Regulation T as a pre-trade\n"
        "compliance engine.\n\n"
        "Tech stack: FastAPI (Python) backend, React/TypeScript frontend,\n"
        "SQLite via SQLAlchemy ORM.\n\n"
        "Default save location: ~/Documents/AI_Study/finra4210/ProjectRefDocs/\n\n"
        "After every answer, ask me if I would like it saved to a Word document.\n"
        "All Word documents must include page numbers and a hyperlinked TOC."
    )
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    r.font.color.rgb = DARK_NAVY
    page_break(doc)

    # ── SECTION 4 ─────────────────────────────────────────────────────────────
    add_heading(doc, "4.  Side-by-Side Comparison", level=1)
    add_table(doc,
        ["Feature", "Claude Code (CLI)", "Claude.ai Web Projects"],
        [
            ["Access method",         "Terminal: claude",                                 "Browser: claude.ai"],
            ["Instructions",          "CLAUDE.md file (project or global)",               "Instructions section in Project settings"],
            ["Persistent memory",     "Memory files at ~/.claude/projects/",              "Conversation history within the Project"],
            ["File read/write",       "Yes — full local filesystem access",               "No — upload only"],
            ["Code execution",        "Yes — runs shell commands, scripts, tests",        "No"],
            ["Session continuity",    "Memory files carry context across sessions",       "Conversation history carries context"],
            ["Project knowledge",     "Reads files on demand from working directory",     "Uploaded files stored in Project Knowledge"],
            ["Best for",              "Engineering tasks, file editing, running the app", "Q&A, writing, analysis, team collaboration"],
            ["Requires installation", "Yes — install via npm or Homebrew",                "No — browser only"],
            ["Offline capable",       "Partially (local files); needs internet for Claude", "No — always requires internet"],
        ],
        col_widths=[Inches(1.7), Inches(2.4), Inches(2.3)]
    )
    page_break(doc)

    # ── SECTION 5 ─────────────────────────────────────────────────────────────
    add_heading(doc, "5.  CLAUDE.md — The Instruction File for Claude Code", level=1)
    add_para(doc,
        "CLAUDE.md is a Markdown file that Claude Code reads automatically when launched "
        "from a directory that contains (or inherits) the file. It is the Claude Code "
        "equivalent of the Instructions section in a Claude.ai Project."
    )

    add_heading(doc, "5.1  Where to Place CLAUDE.md", level=2)
    add_table(doc,
        ["Location", "Scope", "Use Case"],
        [
            ["~/.claude/CLAUDE.md",                          "Global — every directory",             "Instructions you always want regardless of project — response style, document formatting"],
            ["~/Documents/AI_Study/finra4210/CLAUDE.md",     "Project-level — finra4210/ only",      "ArT-specific instructions — save paths, Word doc formatting, prompting preferences"],
            ["Any subdirectory/CLAUDE.md",                   "Subtree — that folder and below",      "Module-specific instructions"],
        ],
        col_widths=[Inches(2.4), Inches(1.8), Inches(3.2)]
    )

    add_heading(doc, "5.2  Current CLAUDE.md Instructions for the ArT Project", level=2)
    add_para(doc, "The following instructions are active in the ArT project CLAUDE.md:", colour=MID_GREY, italic=True)
    for item in [
        "After answering any query, prompt the user whether they would like the answer saved to a Word document. If yes, ask for the file name and the folder path where it should be saved.",
        "Every Word document generated must include page numbers in the footer (Page X of Y, centred).",
        "Every Word document generated must include a Table of Contents with page numbers. The TOC must be hyperlinked so the reader can click an entry and navigate to the relevant page.",
    ]:
        add_bullet(doc, item)

    # ── SAVE ──────────────────────────────────────────────────────────────────
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
