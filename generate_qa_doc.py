"""
Generates the ArT QA & UAT Artefacts Word document.
Run: python3 generate_qa_doc.py
Output: ProjectRefDocs/ArT_QA_UAT_Artefacts.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "ProjectRefDocs", "ArT_QA_UAT_Artefacts.docx")

DARK_NAVY = RGBColor(0x1E, 0x29, 0x3B)
PURPLE    = RGBColor(0x7C, 0x3A, 0xED)
MID_GREY  = RGBColor(0x64, 0x74, 0x8B)
GREEN     = RGBColor(0x16, 0xA3, 0x4A)


def add_page_numbers(doc):
    """Add centred 'Page X of Y' to the footer of every section."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.font.size = Pt(9)
        run.font.color.rgb = MID_GREY

        def append_field(r, instr):
            fc_begin = OxmlElement("w:fldChar")
            fc_begin.set(qn("w:fldCharType"), "begin")
            it = OxmlElement("w:instrText")
            it.set(qn("xml:space"), "preserve")
            it.text = instr
            fc_end = OxmlElement("w:fldChar")
            fc_end.set(qn("w:fldCharType"), "end")
            r._r.append(fc_begin)
            r._r.append(it)
            r._r.append(fc_end)

        t1 = OxmlElement("w:t")
        t1.set(qn("xml:space"), "preserve")
        t1.text = "Page "
        run._r.append(t1)
        append_field(run, " PAGE ")
        t2 = OxmlElement("w:t")
        t2.set(qn("xml:space"), "preserve")
        t2.text = " of "
        run._r.append(t2)
        append_field(run, " NUMPAGES ")


def add_hyperlinked_toc(doc):
    """Insert a TOC field; Word populates hyperlinks and page numbers on open."""
    para = doc.add_paragraph()
    run  = para.add_run()
    fc_begin = OxmlElement("w:fldChar")
    fc_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fc_sep = OxmlElement("w:fldChar")
    fc_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = (
        "[Open in Word, then right-click here and select "
        "'Update Field' → 'Update entire table' to generate the TOC]"
    )
    fc_end = OxmlElement("w:fldChar")
    fc_end.set(qn("w:fldCharType"), "end")
    run._r.append(fc_begin)
    run._r.append(instr)
    run._r.append(fc_sep)
    run._r.append(placeholder)
    run._r.append(fc_end)
    run.font.color.rgb = MID_GREY
    run.italic = True
    run.font.size = Pt(9)
    doc.add_paragraph()


def set_cell_bg(cell, hex_colour):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_colour)
    tcPr.append(shd)


def add_heading(doc, text, level=1, colour=DARK_NAVY):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = colour
    return h


def add_para(doc, text="", bold=False, italic=False, colour=None, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold   = bold
    r.italic = italic
    if colour:
        r.font.color.rgb = colour
    if size:
        r.font.size = Pt(size)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p


def set_header_row(row, texts, bg="1E293B", fg=RGBColor(0xFF, 0xFF, 0xFF)):
    for i, cell in enumerate(row.cells):
        cell.text = texts[i] if i < len(texts) else ""
        set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = fg
                run.font.size = Pt(9)


def add_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_header_row(tbl.rows[0], headers)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = tbl.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for row in tbl.rows:
            for i, cell in enumerate(row.cells):
                if i < len(col_widths):
                    cell.width = col_widths[i]
    return tbl


def page_break(doc):
    doc.add_page_break()


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    add_page_numbers(doc)

    # ── TITLE PAGE ────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("ArT — Automated Review Tool")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = DARK_NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("QA Testing & UAT — Artefacts Guide")
    r2.font.size = Pt(16)
    r2.font.color.rgb = PURPLE

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Version 1.0  |  June 2026").font.color.rgb = MID_GREY
    page_break(doc)

    # ── TABLE OF CONTENTS ─────────────────────────────────────────────────────
    add_heading(doc, "Table of Contents", level=1)
    add_para(doc,
        "After opening in Word, right-click the field below and select "
        "'Update Field' → 'Update entire table' to generate page numbers and hyperlinks.",
        italic=True, colour=MID_GREY, size=9
    )
    add_hyperlinked_toc(doc)
    page_break(doc)

    # ── SECTION 1 ─────────────────────────────────────────────────────────────
    add_heading(doc, "1.  QA Testing Artefacts", level=1)
    add_para(doc,
        "The following artefacts are required to complete QA testing of the ArT product. "
        "All artefacts apply to each QA test cycle and must be baselined before UAT entry.",
        colour=MID_GREY, italic=True
    )
    doc.add_paragraph()
    add_table(doc,
        ["#", "Artefact", "Description", "Scope for ArT"],
        [
            ["1", "Test Strategy",
             "Overall QA approach, objectives, risk-based prioritisation, test types, entry/exit criteria, tools.",
             "Covers all 8 functional areas: Reg T, FINRA 4210, Custom Rules, Override Workflow, AI Scoring, Rules Manager, MI Dashboard, Audit."],
            ["2", "Test Plan",
             "Detailed plan per test cycle — timelines, resource allocation, environments, dependencies, sign-off owners.",
             "Includes backend API testing (FastAPI /docs), frontend UI testing, and integration testing."],
            ["3", "Requirements Traceability Matrix (RTM)",
             "Maps every business requirement (BR-01 to BR-18) to one or more test cases. Ensures no requirement is untested.",
             "Critical for a regulatory product — regulator may ask for evidence that all Rule 4210 requirements are tested."],
            ["4", "Test Cases",
             "Structured test cases with steps, inputs, expected result, pass/fail.",
             "See Section 2 for full coverage breakdown by area."],
            ["5", "Test Data Pack",
             "All input data sets needed to execute test cases, organised by scenario.",
             "Significant portion already built during review sessions (IML scenarios, HBO trades, breach trades)."],
            ["6", "Defect Log / Bug Register",
             "Log of all defects found — severity, priority, status, owner, resolution.",
             "Tracked per test cycle. Severity rated P1 (Critical) to P4 (Minor)."],
            ["7", "Test Execution Report",
             "Results of each test run — pass/fail count, defect count, coverage %.",
             "Produced at the end of each test cycle."],
            ["8", "Regression Test Suite",
             "Subset of critical test cases re-run after every bug fix or code change.",
             "Focused on the rules engine outcome calculations, severity escalation logic, and override workflow."],
            ["9", "Test Summary / Exit Report",
             "Final QA sign-off document — overall quality verdict, outstanding risks, deferred defects.",
             "Required before UAT entry gate is opened."],
            ["10", "Environment Checklist",
             "Confirms test environment is correctly configured — backend running, DB seeded, frontend connected.",
             "Covers alias setup, PATH configuration, uvicorn startup, and seed data verification."],
        ],
        col_widths=[Inches(0.4), Inches(1.6), Inches(2.5), Inches(2.9)]
    )
    page_break(doc)

    # ── SECTION 2 ─────────────────────────────────────────────────────────────
    add_heading(doc, "2.  Test Case Coverage for ArT", level=1)
    add_para(doc,
        "Test cases must cover the following functional areas. Each row represents a testing workstream "
        "with the types of tests required.",
        colour=MID_GREY, italic=True
    )
    doc.add_paragraph()
    add_table(doc,
        ["Area", "Test Case Types"],
        [
            ["Reg T Engine",
             "Happy path (clean pass); each rule breach (REGT-LM-001, SS-001, CA-001, FI-IG-001, FI-HY-001); "
             "borderline/boundary values (coverage at exactly 80% and 85% thresholds); maintenance margin warning (REGT-MM-001)."],
            ["FINRA 4210 Engine",
             "Each IML deficit band: safe harbour (F4210-IML-001), moderate (F4210-IML-002), large (F4210-IML-003), "
             "critical (F4210-IML-004); 90-day freeze short block (F4210-FRZ-001) and buy soft block (F4210-FRZ-002); "
             "portfolio margin sub-$5M (F4210-PM-001)."],
            ["Custom Rules",
             "CUST-001 (analyst notional cap > $5M); CUST-002 (HY bond cap > $10M); toggle rule inactive and verify "
             "engine excludes it; NL Builder creates a rule that fires correctly."],
            ["Severity Escalation",
             "Multiple rules firing simultaneously — verify worst outcome always wins. "
             "e.g. SOFT_BLOCK + HARD_BLOCK → outcome = HARD_BLOCK."],
            ["Override Workflow",
             "Approve: verify trade status → OVERRIDE_APPROVED and audit entry created. "
             "Reject: verify trade status → OVERRIDE_REJECTED. "
             "Escalate: verify status → ESCALATED, trade status unchanged, audit entry created. "
             "Verify only PENDING overrides can be actioned (attempt action on APPROVED → expect error)."],
            ["AI Risk Scoring",
             "Verify each scoring factor contributes correctly: notional size (log scale), IML deficit ratio, "
             "trader seniority (Analyst vs MD), override history (30d), Reg T flag (+10pts). "
             "Verify band thresholds: ≤30=LOW, 31–60=MEDIUM, 61–80=HIGH, >80=CRITICAL."],
            ["Rules Manager",
             "View all rules; filter by REG_T / FINRA_4210 / CUSTOM; toggle active/inactive on a custom rule; "
             "NL Builder: parse plain English → review parsed result → confirm and save → verify rule appears in list and fires on engine."],
            ["MI Dashboard",
             "KPI counts match actual trade/decision data; period filter (7/14/30/90 days) changes results correctly; "
             "charts render with data; NL query returns correct filtered results; suggested queries work."],
            ["Audit Log",
             "Every trade check produces a TRADE_CHECKED entry; every override action produces the correct event type; "
             "filter by event type; filter by date range."],
            ["Negative / Error Testing",
             "Invalid trader ID → 404 error; missing required fields → validation error; "
             "wrong enum values (e.g. invalid direction) → validation error; attempt to action a non-existent override → 404."],
            ["Boundary Testing",
             "IML deficit exactly $1,000 (safe harbour boundary); exactly $25,000 (IML-002 / IML-003 boundary); "
             "exactly $100,000 (IML-003 / IML-004 boundary); notional exactly $5,000,000 (CUST-001 boundary); "
             "notional exactly $10,000,000 (CUST-002 boundary); equity exactly equal to Reg T required (no shortfall)."],
        ],
        col_widths=[Inches(1.8), Inches(5.6)]
    )
    page_break(doc)

    # ── SECTION 3 ─────────────────────────────────────────────────────────────
    add_heading(doc, "3.  Do You Need Separate UAT?", level=1)
    add_para(doc, "Yes — UAT is a separate, mandatory phase for a regulatory product like ArT.")
    doc.add_paragraph()
    add_para(doc, "Why UAT is mandatory for ArT:", bold=True)
    for point in [
        "ArT enforces regulatory rules (FINRA Rule 4210, Regulation T). The business — not just the QA team — must confirm the system behaves correctly in real trading scenarios before it can be used to block or approve live trades.",
        "Regulators may request evidence that business users have validated the control. UAT sign-off is that evidence.",
        "QA tests that the system works correctly. UAT tests that it works correctly for the business — these are different questions.",
        "Real-world edge cases (unusual trade structures, cross-desk scenarios, escalation chains) are best identified by the traders and supervisors who will use the system daily.",
    ]:
        add_bullet(doc, point)
    page_break(doc)

    # ── SECTION 4 ─────────────────────────────────────────────────────────────
    add_heading(doc, "4.  QA vs UAT — Key Differences", level=1)
    doc.add_paragraph()
    add_table(doc,
        ["Dimension", "QA Testing", "UAT"],
        [
            ["Who",
             "QA team (technical testers)",
             "Business users — Traders, Supervisors, Controls Team, Compliance"],
            ["What",
             "Validates the system works correctly against requirements",
             "Validates the system works correctly for the business in real-world conditions"],
            ["Test Cases vs Scenarios",
             "Structured test cases with precise inputs, steps, and expected outputs",
             "End-to-end business scenarios e.g. 'A trader submits a short sell and the supervisor reviews and escalates the override'"],
            ["Data",
             "Synthetic test data designed to hit specific conditions and boundary values",
             "Representative real-world trade flows reflecting actual desk activity and notional sizes"],
            ["Boundary / Negative Tests",
             "Yes — extensively tested with exact threshold values",
             "Minimal — UAT focuses on the golden path and key business journeys"],
            ["Defect Severity",
             "Technical severity: P1 (system down) to P4 (cosmetic)",
             "Business impact severity: blocks a desk, produces a wrong regulatory decision, misleads a supervisor"],
            ["Sign-off Owner",
             "QA Lead signs off test completion",
             "Business owner / Head of Compliance signs off acceptance"],
            ["Regulatory Evidence",
             "Internal QA evidence — not typically shown to regulator",
             "UAT sign-off is the evidence submitted to regulators that the business has validated the control"],
            ["IML / Rules Knowledge Required",
             "QA team needs technical knowledge of the engine to design test cases",
             "UAT participants need business knowledge of FINRA Rule 4210 to assess whether decisions are correct"],
        ],
        col_widths=[Inches(1.8), Inches(2.8), Inches(2.8)]
    )
    page_break(doc)

    # ── SECTION 5 ─────────────────────────────────────────────────────────────
    add_heading(doc, "5.  UAT-Specific Artefacts", level=1)
    add_para(doc,
        "The following artefacts are produced during UAT. Each is distinct from its QA equivalent "
        "in audience, language, and purpose.",
        colour=MID_GREY, italic=True
    )
    doc.add_paragraph()
    add_table(doc,
        ["#", "Artefact", "Description", "How It Differs from QA Equivalent"],
        [
            ["1", "UAT Plan",
             "Overall UAT approach — scope, participants, schedule, entry/exit criteria, environment.",
             "Written in business language; owned by the business, not QA. Defines which desks participate and who the UAT lead is."],
            ["2", "UAT Scenarios",
             "Business journeys that UAT participants execute end-to-end.",
             "Business journeys, not technical test cases. e.g. 'Equity desk Analyst submits a large short position that triggers a 90-day freeze block' — no field-level input/output tables."],
            ["3", "UAT Test Data",
             "Trade inputs used during UAT execution.",
             "Trades reflect real desk activity and realistic notional sizes, not boundary values or edge cases."],
            ["4", "UAT Execution Sign-off Sheet",
             "Record of each scenario executed — who ran it, when, pass/fail, comments.",
             "Each business participant signs off each scenario they ran. QA test execution is recorded by the QA team only."],
            ["5", "UAT Defect Log",
             "Issues raised by business users during UAT.",
             "Written in business language. e.g. 'The AI risk rationale is not clear enough for a supervisor to make a decision' rather than 'field X returns null when Y is zero'."],
            ["6", "Known Issues / Deferred Log",
             "Items the business accepts as known limitations for go-live.",
             "e.g. 'Account state must be entered manually — to be replaced by position system feed in Phase 2.' Business formally accepts these gaps."],
            ["7", "UAT Acceptance Certificate",
             "Formal sign-off document confirming ArT is fit for purpose.",
             "Signed by the business owner (e.g. Head of Compliance or Chief Risk Officer). This is the gate to production. No QA equivalent — QA produces an exit report, not an acceptance certificate."],
        ],
        col_widths=[Inches(0.4), Inches(1.6), Inches(2.5), Inches(2.9)]
    )
    page_break(doc)

    # ── SECTION 6 ─────────────────────────────────────────────────────────────
    add_heading(doc, "6.  Suggested Delivery Sequence", level=1)
    doc.add_paragraph()

    add_para(doc, "Phase 1 — QA Cycle 1 (Functional Testing)", bold=True)
    for item in [
        "Execute all functional test cases across all 8 areas",
        "Log defects in Defect Register",
        "Produce Test Execution Report (Cycle 1)",
    ]:
        add_bullet(doc, item)

    doc.add_paragraph()
    add_para(doc, "Phase 2 — QA Cycle 2 (Regression + Defect Fixes)", bold=True)
    for item in [
        "Re-test all P1/P2 defect fixes",
        "Run full regression test suite",
        "Produce Test Execution Report (Cycle 2)",
        "Produce Test Summary / Exit Report",
        "QA Lead sign-off",
    ]:
        add_bullet(doc, item)

    doc.add_paragraph()
    add_para(doc, "UAT Entry Gate", bold=True)
    add_bullet(doc, "QA Exit Report reviewed and accepted by business owner before UAT begins")

    doc.add_paragraph()
    add_para(doc, "Phase 3 — UAT Execution", bold=True)
    for item in [
        "Business users execute UAT Scenarios",
        "UAT Defects raised and triaged",
        "P1/P2 UAT defects fixed and re-tested",
        "UAT Execution Sign-off Sheets completed",
    ]:
        add_bullet(doc, item)

    doc.add_paragraph()
    add_para(doc, "Phase 4 — UAT Closure", bold=True)
    for item in [
        "Known Issues / Deferred Log agreed with business",
        "UAT Acceptance Certificate signed by business owner",
        "Production release approved",
    ]:
        add_bullet(doc, item)

    doc.add_paragraph()
    doc.add_paragraph()
    add_para(doc, "Summary Sequence Diagram:", bold=True)
    p = doc.add_paragraph()
    r = p.add_run(
        "QA Cycle 1 (Functional)\n"
        "        ↓\n"
        "QA Cycle 2 (Regression + Fixes)\n"
        "        ↓\n"
        "QA Exit Report + Sign-off\n"
        "        ↓\n"
        "UAT Entry Gate\n"
        "        ↓\n"
        "UAT Execution (Business Users)\n"
        "        ↓\n"
        "UAT Defect Fixes + Re-test\n"
        "        ↓\n"
        "UAT Acceptance Certificate\n"
        "        ↓\n"
        "Production Release"
    )
    r.font.name = "Courier New"
    r.font.size = Pt(10)
    r.font.color.rgb = DARK_NAVY

    # ── SAVE ──────────────────────────────────────────────────────────────────
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
